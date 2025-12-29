import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx2pdf import convert
import pythoncom
import os
from flask import send_from_directory
from dotenv import load_dotenv

load_dotenv()

# Caminho onde os currículos serão salvos. Se não estiver no .env, usa o padrão fornecido.
PATH_EXPORT = (
    os.getenv("PATH_EXPORT")
    or r"N:\github\Automacao-Adaptacao-Curriculos\backend\exports\\"
)

# Garante que a pasta de exportação existe
if PATH_EXPORT and not os.path.exists(PATH_EXPORT):
    os.makedirs(PATH_EXPORT, exist_ok=True)

# Constantes globais
COLOR_TITLE = RGBColor(40, 40, 40)
COLOR_PRIMARY = RGBColor(242, 96, 0)
SIZE_SUBTITLE = 15
COLOR_LINK = RGBColor(0, 102, 204)


def add_hyperlink(paragraph, text, url):
    """Adiciona um hyperlink a um parágrafo"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Cor azul para links
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0066CC")
    rPr.append(color)

    # Sublinhado
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def setup_document():
    """Configura o documento com margens e estilos"""
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4 / 2.54)
        section.bottom_margin = Inches(0.8 / 2.54)
        section.left_margin = Inches(1.9 / 2.54)
        section.right_margin = Inches(1.9 / 2.54)

    # Configurar fonte e espaçamento
    style = doc.styles["Normal"]
    style.font.name = "Calibri"

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

    return doc


def add_header(doc, data):
    """Adiciona o cabeçalho do currículo"""
    title = f"{data['nome_candidato']} | {data['objetivo']}"
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.name = "Arial"
    heading.runs[0].font.size = Pt(18)
    heading.runs[0].font.color.rgb = COLOR_TITLE
    return doc


def add_contact_info(doc, data):
    """Adiciona informações de contato"""
    contato = data["informacoes_contato"]

    # Telefone e email
    p = doc.add_paragraph()
    p.add_run("CEL: ").bold = True
    p.add_run(f"{contato['telefone']} | ")
    p.add_run("EMAIL: ").bold = True
    add_hyperlink(p, contato["email"], f'mailto:{contato["email"]}')

    # Endereço
    p = doc.add_paragraph()
    p.add_run("ENDEREÇO: ").bold = True
    p.add_run(contato["endereco"])

    # LinkedIn
    if contato.get("linkedin"):
        print("passo aqui linkedin")
        p = doc.add_paragraph()
        p.add_run("LINKEDIN: ").bold = True
        linkedin_url = (
            contato["linkedin"]
            if contato["linkedin"].startswith("http")
            else f"https://{contato['linkedin']}"
        )
        add_hyperlink(p, contato["linkedin"], linkedin_url)

    # GitHub
    if contato.get("github"):
        print("passo aqui github")
        p = doc.add_paragraph()
        p.add_run("GITHUB: ").bold = True
        add_hyperlink(p, contato["github"], contato["github"])

    # Idiomas
    p = doc.add_paragraph()
    p.add_run("IDIOMAS: ").bold = True
    p.add_run(", ".join(data["idiomas"]))

    return doc


def add_section_heading(doc, title):
    """Adiciona um cabeçalho de seção"""
    heading = doc.add_heading(title, level=2)
    heading.runs[0].font.name = "Arial"
    heading.runs[0].font.size = Pt(SIZE_SUBTITLE)
    heading.runs[0].font.color.rgb = COLOR_PRIMARY
    return doc


def add_summary(doc, data):
    """Adiciona a seção de resumo"""
    add_section_heading(doc, "Resumo")

    p = doc.add_paragraph()
    p.add_run(data["resumo_profissional"])

    return doc


def add_experiences(doc, data):
    """Adiciona a seção de experiências"""
    add_section_heading(doc, "Experiências")

    for experience in data["experiencias_profissionais"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(
            f"{experience['empresa']}, {experience['cargo']} | {experience['tempo_atuacao']}"
        ).bold = True
        p.add_run(f"\n{experience['atividades']}\n")

        is_exist_tecnologias = experience.get("tecnologias")
        if is_exist_tecnologias:
            p.add_run(f"Tecnologias: {experience['tecnologias']};").bold = True
            p.add_run("\n")

    return doc


def add_skills(doc, data):
    """Adiciona a seção de habilidades"""
    add_section_heading(doc, "Competências Técnicas")

    for skill in data["outras_habilidades"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(skill)

    return doc


def add_education(doc, data):
    """Adiciona a seção de formação"""
    add_section_heading(doc, "Formação")

    for formacao in data["formacoes"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(formacao)

    return doc


def convert_to_pdf(docx_file, pdf_file=None):
    """Converte um arquivo .docx para PDF"""
    try:
        if pdf_file is None:
            pdf_file = docx_file.replace(".docx", ".pdf")

        # Inicializar COM
        pythoncom.CoInitialize()

        # Converte o arquivo
        convert(docx_file, pdf_file)
        print(f"Arquivo convertido para PDF: {pdf_file}")
        return pdf_file
    except Exception as e:
        print(f"Erro ao converter para PDF: {e}")
        return None
    finally:
        # Limpar COM
        pythoncom.CoUninitialize()


def load_resume_data(json_file_path):
    """Carrega os dados do currículo a partir de um arquivo JSON"""
    try:
        with open(json_file_path, "r", encoding="utf-8") as file:
            return json.load(file)
    except FileNotFoundError:
        print(f"Arquivo {json_file_path} não encontrado!")
        return None
    except json.JSONDecodeError:
        print(f"Erro ao decodificar o JSON do arquivo {json_file_path}")
        return None


def get_datetime_now():
    return datetime.now().strftime("%d-%m-%Y_%H-%M-%S")


def create_document_pdf_and_docx(data: dict):
    print("Criando currículo...")
    """Função principal que cria o currículo completo a partir de um arquivo JSON"""
    # Configurar documento
    doc = setup_document()

    # Adicionar seções
    doc = add_header(doc, data)
    doc = add_contact_info(doc, data)
    doc = add_summary(doc, data)
    doc = add_experiences(doc, data)
    doc = add_skills(doc, data)
    doc = add_education(doc, data)

    # Salvar documento
    nome_arquivo = data["nome_candidato"].lower().replace(" ", "_")
    company_name = data["nome_da_empresa_canditatura"].lower().replace(" ", "_")

    filename_docx = f"curriculo_{nome_arquivo}_{company_name}.docx"
    docx_filename = PATH_EXPORT + filename_docx
    print(docx_filename)
    doc.save(docx_filename)
    print("Currículo criado com sucesso!")

    # Converter para PDF
    pdf_filename = convert_to_pdf(docx_filename)
    if pdf_filename:
        print(f"Currículo também salvo em PDF: {pdf_filename}")

    send_from_directory(PATH_EXPORT, filename_docx, as_attachment=True)
    # send_from_directory(PATH_EXPORT, pdf_filename)
    return docx_filename, pdf_filename
