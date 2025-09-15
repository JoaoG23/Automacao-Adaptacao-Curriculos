from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx2pdf import convert
from time import sleep

# Constantes globais
COLOR_TITLE = RGBColor(40, 40, 40)
COLOR_PRIMARY = RGBColor(242, 96, 0)
SIZE_SUBTITLE = 15
COLOR_LINK = RGBColor(0, 102, 204)

def add_hyperlink(paragraph, text, url):
    """Adiciona um hyperlink a um parágrafo"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Cor azul para links
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066CC')
    rPr.append(color)
    
    # Sublinhado
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def setup_document():
    """Configura o documento com margens e estilos"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4 / 2.54)
        section.bottom_margin = Inches(0.8 / 2.54)
        section.left_margin = Inches(1.9 / 2.54)
        section.right_margin = Inches(1.9 / 2.54)
    
    # Configurar fonte e espaçamento
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    
    return doc

def add_header(doc):
    """Adiciona o cabeçalho do currículo"""
    heading = doc.add_heading('JOÃO GUILHERME | Desenvolvedor Fullstack JS', level=1)
    heading.runs[0].font.name = 'Arial'
    heading.runs[0].font.size = Pt(18)
    heading.runs[0].font.color.rgb = COLOR_TITLE
    return doc

def add_contact_info(doc):
    """Adiciona informações de contato"""
    # Telefone e email
    p = doc.add_paragraph()
    p.add_run('CEL: ').bold = True
    p.add_run('(31) 99621-6938 | ')
    p.add_run('EMAIL: ').bold = True
    add_hyperlink(p, 'jguilhermeempresarial@outlook.com', 'mailto:jguilhermeempresarial@outlook.com')
    
    # Endereço
    p = doc.add_paragraph()
    p.add_run('ENDEREÇO: ').bold = True
    p.add_run('Ribeirão das Neves, Minas Gerais, BR')
    
    # LinkedIn
    p = doc.add_paragraph()
    p.add_run('LINKEDIN: ').bold = True
    add_hyperlink(p, 'www.linkedin.com/in/joaog123', 'https://www.linkedin.com/in/joaog123')
    
    # Portfólio
    p = doc.add_paragraph()
    p.add_run('PORTFÓLIO: ').bold = True
    add_hyperlink(p, 'https://joao-guilherme-portifolio.vercel.app', 'https://joao-guilherme-portifolio.vercel.app')
    
    # GitHub
    p = doc.add_paragraph()
    p.add_run('GITHUB: ').bold = True
    add_hyperlink(p, 'https://github.com/JoaoG23', 'https://github.com/JoaoG23')
    
    # Idiomas
    p = doc.add_paragraph()
    p.add_run('IDIOMAS: ').bold = True
    p.add_run('Inglês (Intermediário), Espanhol (Intermediário)')
    
    return doc

def add_section_heading(doc, title):
    """Adiciona um cabeçalho de seção"""
    heading = doc.add_heading(title, level=2)
    heading.runs[0].font.name = 'Arial'
    heading.runs[0].font.size = Pt(SIZE_SUBTITLE)
    heading.runs[0].font.color.rgb = COLOR_PRIMARY
    return doc

def add_summary(doc):
    """Adiciona a seção de resumo"""
    add_section_heading(doc, 'Resumo')
    
    p = doc.add_paragraph()
    p.add_run('Tenho 4 anos de experiência em desenvolvimento Fullstack, especializado em websites, '
              'automações, integrações, sistemas comerciais, e sou proficiente em tecnologias como '
              'Node.js, Typescript, React.js, entre outras.')
    
    return doc

def add_experiences(doc):
    """Adiciona a seção de experiências"""
    add_section_heading(doc, 'Experiências')
    
    # Mirror Tecnologia
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Mirror Tecnologia, Desenvolvedor Fullstack | Fev de 2025 - Atualmente').bold = True
    p.add_run('\nAtuo em soluções de e-commerce e logística, com especialização em integrações entre marketplaces '
              '(Mercado Livre, Amazon, Magalu) e ERPs (Bling, Tiny, SAP). Participei de projetos estratégicos para '
              'grandes empresas, como a Mondial Eletrodomésticos, solucionando desafios técnicos e operacionais. '
              'Sou coautor de um SaaS analítico voltado para vendas no Mercado Livre, responsável por entregar '
              'insights valiosos como ticket médio, receita consolidada e tendências de mercado.\n')
    p.add_run('Tecnologias: NodeJS, NestJS, Python, MySQL, SQLAlchemy, Flask, Webhook, Oauth2;').bold = True
    p.add_run('\n')
    
    # Autônomo 2024-2025
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Autônomo, Desenvolvedor Fullstack | Ago de 2024 - Fev de 2025').bold = True
    p.add_run('\nCriei sistemas financeiros e contábeis que otimizaram processos. Além disso, implementei automações '
              'para extrair dados de NFEs em PDFs, permitindo que balanços que antes levavam dias para serem concluídos '
              'fossem realizados em apenas algumas horas, reduzindo o trabalho do cliente.\n')
    p.add_run('Tecnologias: ReactJS, NodeJS, ExpressJS, Typescript, PostgresSQL, Docker, Python, RPA, Selenium;').bold = True
    p.add_run('\n')
    
    # XC Soluções Tecnológicas
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('XC Soluções Tecnológicas, Desenvolvedor Fullstack Web | Dez de 2022 - Mai de 2024').bold = True
    p.add_run('\nOtimizei a gestão de acesso para caminhoneiros e veículos, criando, estruturando e documentando o '
              'Sistema de Controle de Acesso para Docas e Portarias, do início ao fim, com atuação no Back-end e '
              'Front-end. Integrei uma balança rodoviária ao Sistema de Gestão de Doca e Portaria, automatizando o '
              'check-in de veículos, o que melhorou significativamente o processo logístico e eliminou filas de '
              'caminhões.\n')
    p.add_run('Tecnologias: ReactJS, Redux, Typescript, Bootstrap, Material UI, ExpressJS, SQL Server, PostgresSQL, Oracle SQL, '
              'PrismaORM, NodeJS, Git, Github, Docker;').bold = True
    p.add_run('\n')
    
    # Autônomo 2022
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Autônomo, Desenvolvedor Freelancer Web | Mai de 2022 - Nov de 2022').bold = True
    p.add_run('\nAtuei como freelancer no desenvolvimento de um MVP de aplicativo para academia, envolvendo a criação '
              'do Back-end e Front-end. Também participei do levantamento de requisitos para garantir soluções '
              'alinhadas às necessidades do projeto.\n')
    p.add_run('Tecnologias: ReactJS, NodeJS, ExpressJS, Typescript, SequelizeORM;').bold = True
    p.add_run('\n')
    
    # AGS INN
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AGS INN, Técnico de Suporte em Desenvolvimento TI e Software | Jul de 2021 - Abr de 2022').bold = True
    p.add_run('\nSuporte técnico aos usuários, implantação de sistemas e criação de interfaces Front-end para sistemas.\n')
    p.add_run('Tecnologias: HTML, CSS, Bootstrap, Materialize, Java, Angular, Ionic, MySQL, PostgresSQL, SOUP, REST, Git;').bold = True
    p.add_run('\n')
    
    return doc

def add_skills(doc):
    """Adiciona a seção de habilidades"""
    add_section_heading(doc, 'Outras Habilidades')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Backend: NestJS, Pandas, Flask, Java, Spring Boot, SQL, MongoDB, Redis;')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frontend: TailwindCSS, Styled-Components, NextJS;')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Testes: JEST, Testing-library, Cypress, JUnit, Mockito, Selenium;')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Deploys: CI/CD (GitHub Actions), Docker (Kubernetes), AWS (RDS, IAM, VPC, EC2) OCI (OCI, Oracle Database, Oracle Cloud Infrastructure), Mensageria (Apache Kafka);')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Outros: Clean Code, Design patterns, SOLID, Microsserviços, Engenharia de prompt;')
    
    return doc

def add_education(doc):
    """Adiciona a seção de formação"""
    add_section_heading(doc, 'Formação')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Técnico em Informática, Senac Minas, 02/2020 - 01/2022;')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Graduação, Análise e Desenvolvimento de Sistemas - 02/2024 - Cursando atualmente;')
    
    return doc

def convert_to_pdf(docx_file, pdf_file=None):
    """Converte um arquivo .docx para PDF"""
    try:
        if pdf_file is None:
            # Se não especificado, usa o mesmo nome do arquivo .docx mas com extensão .pdf
            pdf_file = docx_file.replace('.docx', '.pdf')
        
        # Converte o arquivo
        convert(docx_file, pdf_file)
        print(f"Arquivo convertido para PDF: {pdf_file}")
        return pdf_file
    except Exception as e:
        print(f"Erro ao converter para PDF: {e}")
        return None

def create_curriculum():
    """Função principal que cria o currículo completo"""
    # Configurar documento
    doc = setup_document()
    
    # Adicionar seções
    doc = add_header(doc)
    doc = add_contact_info(doc)
    doc = add_summary(doc)
    doc = add_experiences(doc)
    doc = add_skills(doc)
    doc = add_education(doc)
    
    # Salvar documento
    docx_filename = 'curriculo_joao_guilherme.docx'
    doc.save(docx_filename)
    print("Currículo criado com sucesso!")
    
    # Converter para PDF
    sleep(7)
    pdf_filename = convert_to_pdf(docx_filename)
    if pdf_filename:
        print(f"Currículo também salvo em PDF: {pdf_filename}")

if __name__ == "__main__":
    create_curriculum()